import io
import re
import logging
from typing import Tuple

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

import PyPDF2
import docx

class TextExtractor:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def extract_text(self, file_content: bytes, filename: str) -> Tuple[str, str]:
        """Extract text from file content"""
        file_ext = filename.lower().split('.')[-1]
        
        try:
            if file_ext == 'pdf':
                return self.extract_from_pdf(file_content)
            elif file_ext in ['docx', 'doc']:
                return self.extract_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            self.logger.error(f"Text extraction failed: {str(e)}")
            raise
    
    def extract_from_pdf(self, file_content: bytes) -> Tuple[str, str]:
        """Extract text from PDF using multiple methods"""
        
        self.logger.info(f"Starting PDF extraction, file size: {len(file_content)} bytes")
        
        # Method 1: Try pdfplumber
        if HAS_PDFPLUMBER:
            try:
                self.logger.info("Trying pdfplumber extraction...")
                pdf_file = io.BytesIO(file_content)
                text = ""
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if len(text.strip()) > 50:
                    self.logger.info(f"pdfplumber successful: {len(text)} chars")
                    return self.clean_text(text), "pdfplumber_extraction"
                else:
                    self.logger.warning("pdfplumber found no text, trying OCR...")
            except Exception as e:
                self.logger.warning(f"pdfplumber failed: {e}")
        
        # Method 2: Try PyPDF2
        try:
            self.logger.info("Trying PyPDF2 extraction...")
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file, strict=False)
            
            text = ""
            for page in pdf_reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception:
                    continue
            
            if len(text.strip()) > 50:
                self.logger.info(f"PyPDF2 successful: {len(text)} chars")
                return self.clean_text(text), "pypdf2_extraction"
        except Exception as e:
            self.logger.warning(f"PyPDF2 failed: {e}")
        
        # Method 3: OCR as last resort
        if HAS_OCR:
            return self.extract_pdf_with_ocr(file_content)
        
        raise Exception("All PDF extraction methods failed. Install OCR dependencies for image-based PDFs.")
    
    def extract_pdf_with_ocr(self, file_content: bytes) -> Tuple[str, str]:
        """Extract text from PDF using OCR"""
        try:
            self.logger.info("Attempting OCR extraction...")
            images = convert_from_bytes(file_content)
            text = ""
            
            for i, image in enumerate(images):
                self.logger.info(f"OCR processing page {i+1}/{len(images)}")
                page_text = pytesseract.image_to_string(image)
                text += page_text + "\n"
            
            self.logger.info(f"OCR successful: {len(text)} chars extracted")
            return self.clean_text(text), "ocr_extraction"
        
        except Exception as e:
            self.logger.error(f"OCR extraction failed: {e}")
            raise Exception(f"OCR extraction failed: {str(e)}")
    
    def extract_from_docx(self, file_content: bytes) -> Tuple[str, str]:
        """Extract text from DOCX"""
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return self.clean_text(text), "docx_extraction"
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        return text.strip()
    
    def extract_contact_patterns(self, text: str) -> dict:
        """Extract contact information using regex patterns"""
        contact_info = {}
        
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text, re.IGNORECASE)
        if emails:
            contact_info['email'] = emails[0]
        
        phone_pattern = r'(\+?1?[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact_info['phone'] = ''.join(phones[0])
        
        linkedin_pattern = r'linkedin\.com/in/[\w\-]+'
        linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_matches:
            contact_info['linkedin'] = linkedin_matches[0]
        
        github_pattern = r'github\.com/[\w\-]+'
        github_matches = re.findall(github_pattern, text, re.IGNORECASE)
        if github_matches:
            contact_info['github'] = github_matches[0]
        
        return contact_info